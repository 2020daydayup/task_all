import docx

class AddDocx():
    "新增word类"
    def __init__(self):
        self.doc = docx.Document()

    def add_title(self, text='这是一个默认的标题', level=0):
        "添加标题"
        self.doc.add_heading(text,level)

    def add_paragraph(self, text='这是一个默认的段落', style=None):
        "添加段落"
        self.doc.add_paragraph(text, style)

    def add_pic(self,path, width=None, height=None):
        "添加图片"
        self.doc.add_picture(path, width=width, height=height)

    def save_docx(self,fileName):
        "保存word文件"
        self.doc.save(fileName)

if __name__ == '__main__':
    doc = AddDocx()
    for i in range(10):
        doc.add_title(f"这是一个{i}级标题",i)
    doc.add_paragraph()
    doc.add_pic("py.jpg")
    doc.save_docx("demo.docx")
